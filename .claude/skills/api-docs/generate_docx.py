"""Render an enterprise-style API reference .docx from a JSON spec.

Used by the api-docs Claude Code skill. Reads a JSON spec on stdin describing
one endpoint and writes a Word document to the path given as argv[1]. Styled to
match a typical enterprise API spec document: built-in Word Title/Heading 1/
Heading 2 styles, light-gray table header rows, plain table bodies, Letter page
size with 1-inch margins, an optional small org logo above the title.

Spec schema (see SKILL.md for how the skill builds this from static analysis):

{
  "title":       "GET Products List API",         # required
  "endpoint":    "GET /api/products/",            # required
  "description": "<long-form description>",       # required

  "auth_permissions": [                            # rendered as Requirement|Detail table
    {"requirement": "Authentication", "detail": "..."},
    {"requirement": "Permissions",    "detail": "..."}
  ],

  "path_params":  [                                # 4-col table
    {"name": "...", "type": "...", "required": "Yes", "description": "..."}
  ],

  "query_params": [                                # 4 or 5 col table (Format col optional)
    {"name": "...", "type": "...", "required": "No",
     "format": "...", "description": "..."}
  ],
  "query_param_subsections": [                     # optional Heading 2 sub-blocks
    {"heading": "Date Range Rules", "paragraphs": ["...", "..."]}
  ],

  "request_body": "None for GET endpoints.",       # str or list of field rows
  # OR
  "request_body": [
    {"name": "...", "type": "...", "required": "Yes", "description": "..."}
  ],

  "responses": [                                   # one entry per variant
    {
      "heading":        "Response — Foreclosure",
      "status":         "HTTP 200 OK",
      "json":           "{\\n  \\"key\\": value\\n}",
      "fields_heading": "Foreclosure Response Fields",
      "fields":         [{"name": "...", "description": "..."}],
      "trailing_notes": ["Note: ..."]
    }
  ],

  "errors": [                                      # HTTP|Condition table
    {"http": "400", "condition": "Validation error..."}
  ]
}
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

HEADER_FILL = 'F3F3F3'
BORDER_GRAY = 'BFBFBF'

SKILL_DIR = Path(__file__).resolve().parent
LOGO = SKILL_DIR / 'shopstack-logo.png'


# ----- low-level OOXML helpers ----------------------------------------------

def _shade(cell, fill_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_borders(cell):
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), BORDER_GRAY)
        borders.append(b)
    cell._tc.get_or_add_tcPr().append(borders)


# ----- content helpers -------------------------------------------------------

def _para(doc, text, *, style=None, italic=False, mono=False, color=None,
          size=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if text == '':
        return p
    run = p.add_run(text)
    if mono:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    if size is not None:
        run.font.size = Pt(size)
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color
    return p


def _h1(doc, text):
    return doc.add_paragraph(text, style='Heading 1')


def _h2(doc, text):
    return doc.add_paragraph(text, style='Heading 2')


def _table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.autofit = True
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        _shade(cell, HEADER_FILL)
        _cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val) if val is not None else '')
            run.font.size = Pt(10)
            _cell_borders(cell)
    return tbl


def _code_block(doc, text):
    """Render a JSON / code block as monospaced paragraphs with light shading."""
    for line in text.splitlines() or [text]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), HEADER_FILL)
        p._p.get_or_add_pPr().append(shd)
    # trailing blank line for breathing room
    doc.add_paragraph()


# ----- section renderers -----------------------------------------------------

def _render_auth(doc, entries):
    _h1(doc, 'Authentication & Permissions')
    if not entries:
        _para(doc, 'Not specified.', italic=True)
        return
    _table(doc, ['Requirement', 'Detail'],
           [[e['requirement'], e['detail']] for e in entries])


def _render_path_params(doc, params):
    _h1(doc, 'Path Parameters')
    if not params:
        _para(doc, 'None.', italic=True)
        return
    _table(doc, ['Parameter', 'Type', 'Required', 'Description'],
           [[p['name'], p['type'], p.get('required', 'No'),
             p.get('description', '')] for p in params])


def _render_query_params(doc, params, subsections):
    _h1(doc, 'Query Parameters')
    if not params:
        _para(doc, 'None.', italic=True)
    else:
        use_format = any(p.get('format') for p in params)
        if use_format:
            _table(doc,
                   ['Parameter', 'Type', 'Required', 'Format', 'Description'],
                   [[p['name'], p['type'], p.get('required', 'No'),
                     p.get('format', ''), p.get('description', '')]
                    for p in params])
        else:
            _table(doc,
                   ['Parameter', 'Type', 'Required', 'Description'],
                   [[p['name'], p['type'], p.get('required', 'No'),
                     p.get('description', '')] for p in params])
    for sub in subsections or []:
        _h2(doc, sub['heading'])
        for line in sub.get('paragraphs', []):
            _para(doc, line)


def _render_request_body(doc, body):
    _h1(doc, 'Request Body')
    if body is None or body == '' or body == []:
        _para(doc, 'None.', italic=True)
        return
    if isinstance(body, str):
        _para(doc, body)
        return
    # list of field dicts
    _table(doc, ['Field', 'Type', 'Required', 'Description'],
           [[f['name'], f.get('type', ''), f.get('required', 'No'),
             f.get('description', f.get('notes', ''))] for f in body])


def _render_responses(doc, responses):
    for resp in responses or []:
        _h1(doc, resp['heading'])
        if resp.get('status'):
            _para(doc, resp['status'])
        if resp.get('json'):
            _code_block(doc, resp['json'])
        if resp.get('fields'):
            _h2(doc, resp.get('fields_heading', 'Response Fields'))
            _table(doc, ['Field', 'Description'],
                   [[f['name'], f['description']] for f in resp['fields']])
        for note in resp.get('trailing_notes', []):
            _para(doc, note)


def _render_errors(doc, errors):
    _h1(doc, 'Error Responses')
    if not errors:
        _para(doc, 'None documented.', italic=True)
        return
    _table(doc, ['HTTP', 'Condition'],
           [[e['http'], e['condition']] for e in errors])


# ----- page setup -----------------------------------------------------------

def _setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def _logo_strip(doc):
    """Org logo, centered, above the title."""
    if not LOGO.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(LOGO), width=Cm(9.5))


# ----- main entry -----------------------------------------------------------

def render(spec, out_path):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    _setup_page(doc)
    _logo_strip(doc)

    # Title
    doc.add_paragraph(spec['title'], style='Title')

    # Endpoint
    _h1(doc, 'Endpoint')
    _para(doc, spec['endpoint'], mono=True)

    # Description
    _h1(doc, 'Description')
    _para(doc, spec['description'])

    # Authentication
    _render_auth(doc, spec.get('auth_permissions', []))

    # Path params
    _render_path_params(doc, spec.get('path_params', []))

    # Query params (+ sub-sections)
    _render_query_params(doc, spec.get('query_params', []),
                         spec.get('query_param_subsections', []))

    # Request body
    _render_request_body(doc, spec.get('request_body'))

    # Responses (one or many)
    _render_responses(doc, spec.get('responses', []))

    # Errors
    _render_errors(doc, spec.get('errors', []))

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main():
    if len(sys.argv) != 2:
        print('usage: generate_docx.py <output.docx>  (spec JSON on stdin)',
              file=sys.stderr)
        sys.exit(2)
    spec = json.load(sys.stdin)
    render(spec, sys.argv[1])
    print(f'wrote: {sys.argv[1]}')


if __name__ == '__main__':
    main()
